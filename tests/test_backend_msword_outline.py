# SPDX-FileCopyrightText: The Docling Contributors
# SPDX-License-Identifier: MIT

"""Tests for Word heading detection via the ``w:outlineLvl`` style property.

Kept separate from ``test_backend_msword.py`` so that file stays under the
repository's per-file line limit.
"""

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docling.datamodel.base_models import InputFormat
from docling.document_converter import DocumentConverter


def _set_outline_level(style, outline_lvl: int):
    """Pin an explicit ``w:outlineLvl`` onto an existing paragraph style."""
    lvl = OxmlElement("w:outlineLvl")
    lvl.set(qn("w:val"), str(outline_lvl))
    style.element.get_or_add_pPr().append(lvl)
    return style


def _add_style_with_outline_level(doc, style_id: str, name: str, outline_lvl: int):
    """Register a paragraph style carrying an explicit ``w:outlineLvl``."""
    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.element.set(qn("w:styleId"), style_id)
    return _set_outline_level(style, outline_lvl)


def _markdown(doc, tmp_path, name: str) -> str:
    docx_path = tmp_path / f"{name}.docx"
    doc.save(str(docx_path))
    converter = DocumentConverter(allowed_formats=[InputFormat.DOCX])
    return converter.convert(docx_path).document.export_to_markdown()


def test_localized_heading_styles_are_detected_via_outline_level(tmp_path):
    """A heading style named in another language must still parse as a heading.

    LibreOffice writes the *localized* style id and name when it converts a
    legacy ``.doc`` (a Czech document arrives as ``w:styleId="Nadpis1"`` with
    ``w:name w:val="Nadpis [1]"`` and no ``basedOn``), but it always records the
    level in ``w:outlineLvl``. Detection used to key solely off the substring
    "heading", so every heading in such a document was emitted as body text and
    the whole section hierarchy was lost.
    """

    def build(localized: bool) -> str:
        doc = Document()
        if localized:
            first = _add_style_with_outline_level(doc, "Nadpis1", "Nadpis [1]", 0)
            second = _add_style_with_outline_level(doc, "Nadpis2", "Nadpis [2]", 1)
        else:
            first = doc.styles["Heading 1"]
            second = doc.styles["Heading 2"]

        doc.add_paragraph("Uvod do problematiky").style = first
        doc.add_paragraph("Body text under the first heading.")
        doc.add_paragraph("Podrobnosti").style = second
        doc.add_paragraph("More body text.")

        name = "localized" if localized else "english"
        return _markdown(doc, tmp_path, name)

    localized = build(localized=True)

    # Naming the styles in Czech may not change the structure of the output.
    assert localized == build(localized=False)

    lines = [line for line in localized.splitlines() if line.strip()]
    assert lines == [
        "## Uvod do problematiky",
        "Body text under the first heading.",
        "### Podrobnosti",
        "More body text.",
    ]


def test_body_text_outline_level_is_not_promoted_to_a_heading(tmp_path):
    """``w:outlineLvl`` 9 means "body text" and must not create a heading.

    Only 0-8 denote heading levels 1-9 in OOXML, so a style pinned to 9 has to
    keep flowing through the regular text path.
    """
    doc = Document()
    style = _add_style_with_outline_level(doc, "BodyPinned", "Body Pinned", 9)
    doc.add_paragraph("Not a heading at all").style = style

    markdown = _markdown(doc, tmp_path, "body_level")

    assert "Not a heading at all" in markdown
    assert "#" not in markdown


def test_title_style_is_not_reclassified_as_a_heading(tmp_path):
    """A ``Title`` style keeps its own branch even when it carries an outline level.

    Word's built-in ``Title`` style defines no ``w:outlineLvl``, but a document
    may still pin one onto it; that must not turn the title into a heading.
    """

    def build(pinned: bool) -> str:
        doc = Document()
        if pinned:
            _set_outline_level(doc.styles["Title"], 0)
        doc.add_paragraph("The document title").style = doc.styles["Title"]
        doc.add_paragraph("Body text.")
        return _markdown(doc, tmp_path, "title_pinned" if pinned else "title_plain")

    assert build(pinned=True) == build(pinned=False)


def test_heading_style_with_the_body_text_sentinel_falls_back_to_its_name(tmp_path):
    """``w:outlineLvl`` 9 does not yield a level-10 heading on a heading style.

    The sentinel is outside the 1-9 range, so the level has to keep coming from
    the style name.
    """

    def build(pinned: bool) -> str:
        doc = Document()
        style = doc.styles.add_style("Custom Heading 3", WD_STYLE_TYPE.PARAGRAPH)
        style.element.set(qn("w:styleId"), "Heading3Alt")
        if pinned:
            _set_outline_level(style, 9)
        doc.add_paragraph("Section title").style = style
        return _markdown(doc, tmp_path, "sentinel" if pinned else "unpinned")

    pinned = build(pinned=True)
    assert pinned == build(pinned=False)
    assert pinned.strip().startswith("#")
