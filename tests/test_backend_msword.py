import logging
import os
from pathlib import Path
from types import SimpleNamespace

import pytest
from docling_core.types.doc import DocItemLabel, GroupItem
from lxml import etree

import docling.backend.msword_backend as msword_backend_module
from docling.backend.docx.drawingml.utils import get_libreoffice_cmd
from docling.backend.msword_backend import MsWordDocumentBackend
from docling.datamodel.base_models import InputFormat
from docling.datamodel.document import (
    ConversionResult,
    DoclingDocument,
    InputDocument,
    SectionHeaderItem,
    TextItem,
)
from docling.document_converter import DocumentConverter

from .test_data_gen_flag import GEN_TEST_DATA
from .verify_utils import verify_document, verify_export

_log = logging.getLogger(__name__)

GENERATE = GEN_TEST_DATA
IS_CI = bool(os.getenv("CI"))


@pytest.fixture(scope="module")
def docx_paths() -> list[Path]:
    # Define the directory you want to search
    directory = Path("./tests/data/docx/")

    # List all docx files in the directory and its subdirectories
    docx_files = sorted(directory.rglob("*.docx"))

    return docx_files


def get_converter():
    converter = DocumentConverter(allowed_formats=[InputFormat.DOCX])

    return converter


@pytest.fixture(scope="module")
def backend(docx_paths) -> MsWordDocumentBackend:
    docx_path = docx_paths[0]
    in_doc = InputDocument(
        path_or_stream=docx_path,
        format=InputFormat.DOCX,
        backend=MsWordDocumentBackend,
    )
    return in_doc._backend


@pytest.fixture(scope="module")
def documents(docx_paths) -> list[tuple[Path, DoclingDocument]]:
    documents: list[dict[Path, DoclingDocument]] = []

    converter = get_converter()

    for docx_path in docx_paths:
        _log.debug(f"converting {docx_path}")

        gt_path = (
            docx_path.parent.parent / "groundtruth" / "docling_v2" / docx_path.name
        )

        conv_result: ConversionResult = converter.convert(docx_path)

        doc: DoclingDocument = conv_result.document

        assert doc, f"Failed to convert document from file {gt_path}"
        documents.append((gt_path, doc))

    return documents


def _test_e2e_docx_conversions_impl(docx_paths: list[tuple[Path, DoclingDocument]]):
    has_libreoffice = False
    try:
        cmd = get_libreoffice_cmd(raise_if_unavailable=True)
        if cmd is not None:
            has_libreoffice = True
    except Exception:
        pass

    for docx_path, doc in docx_paths:
        if not IS_CI and not has_libreoffice and docx_path.name == "drawingml.docx":
            print(f"Skipping {docx_path} because no Libreoffice is installed.")
            continue

        pred_md: str = doc.export_to_markdown()
        assert verify_export(pred_md, str(docx_path) + ".md", generate=GENERATE), (
            f"export to markdown failed on {docx_path}"
        )

        pred_itxt: str = doc._export_to_indented_text(
            max_text_len=70, explicit_tables=False
        )
        assert verify_export(pred_itxt, str(docx_path) + ".itxt", generate=GENERATE), (
            f"export to indented-text failed on {docx_path}"
        )

        assert verify_document(doc, str(docx_path) + ".json", generate=GENERATE), (
            f"DoclingDocument verification failed on {docx_path}"
        )

        if docx_path.name in {"word_tables.docx", "docx_rich_cells.docx"}:
            pred_html: str = doc.export_to_html()
            assert verify_export(
                pred_text=pred_html,
                gtfile=str(docx_path) + ".html",
                generate=GENERATE,
            ), f"export to html failed on {docx_path}"


flaky_file = "textbox.docx"


def test_e2e_docx_conversions(documents):
    target = [item for item in documents if item[0].name != flaky_file]
    _test_e2e_docx_conversions_impl(target)


@pytest.mark.xfail(strict=False)
def test_textbox_conversion(documents):
    target = [item for item in documents if item[0].name == flaky_file]
    _test_e2e_docx_conversions_impl(target)


@pytest.mark.xfail(strict=False)
def test_textbox_extraction(documents):
    name = "textbox.docx"
    doc = next(item[1] for item in documents if item[0].name == name)

    # Verify if a particular textbox content is extracted
    textbox_found = False
    for item, _ in doc.iterate_items():
        if item.text[:30] == """Suggested Reportable Symptoms:""":
            textbox_found = True
    assert textbox_found


def test_heading_levels(documents):
    name = "word_sample.docx"
    doc = next(item[1] for item in documents if item[0].name == name)

    found_lvl_1 = found_lvl_2 = False
    for item, _ in doc.iterate_items():
        if isinstance(item, SectionHeaderItem):
            if item.text == "Let\u2019s swim!":
                found_lvl_1 = True
                assert item.level == 1
            elif item.text == "Let\u2019s eat":
                found_lvl_2 = True
                assert item.level == 2
    assert found_lvl_1 and found_lvl_2


def test_text_after_image_anchors(documents):
    """Test to analyse whether text gets parsed after image anchors."""

    name = "word_image_anchors.docx"
    doc = next(item[1] for item in documents if item[0].name == name)

    found_text_after_anchor_1 = found_text_after_anchor_2 = (
        found_text_after_anchor_3
    ) = found_text_after_anchor_4 = False
    for item, _ in doc.iterate_items():
        if isinstance(item, TextItem):
            if item.text == "This is test 1":
                found_text_after_anchor_1 = True
            elif item.text == "0:08\nCorrect, he is not.":
                found_text_after_anchor_2 = True
            elif item.text == "This is test 2":
                found_text_after_anchor_3 = True
            elif item.text == "0:16\nYeah, exactly.":
                found_text_after_anchor_4 = True

    assert (
        found_text_after_anchor_1
        and found_text_after_anchor_2
        and found_text_after_anchor_3
        and found_text_after_anchor_4
    )


def test_is_rich_table_cell(docx_paths):
    """Test the function is_rich_table_cell."""

    name = "docx_rich_cells.docx"
    path = next(item for item in docx_paths if item.name == name)

    in_doc = InputDocument(
        path_or_stream=path,
        format=InputFormat.DOCX,
        backend=MsWordDocumentBackend,
        filename=name,
    )
    backend = MsWordDocumentBackend(
        in_doc=in_doc,
        path_or_stream=path,
    )

    gt_cells: list[bool] = []
    # table: Table with rich cells
    gt_cells.extend([False, False, True, True, True, True, True, False])
    # table: Table with nested table
    gt_cells.extend([False, False, False, True, True, True])
    # table: Table with pictures
    gt_cells.extend([False, False, False, True, True, False])
    # table: Lists with same numId in different cells
    gt_cells.extend([True, True])
    # table: Lists with different numIds in different cells
    gt_cells.extend([True, True])
    # table: Multiple columns with lists
    gt_cells.extend([True, True, True, True])
    # table: Mixed content - list and regular text in different cells
    gt_cells.extend([True, False])
    gt_it = iter(gt_cells)

    for idx_t, table in enumerate(backend.docx_obj.tables):
        for idx_r, row in enumerate(table.rows):
            for idx_c, cell in enumerate(row.cells):
                assert next(gt_it) == backend._is_rich_table_cell(cell), (
                    f"Wrong cell type in table {idx_t}, row {idx_r}, col {idx_c} "
                    f"with text: {cell.text}"
                )


def test_add_header_footer(documents):
    """Test the funciton _add_header_footer."""

    name = "unit_test_formatting.docx"
    doc = next(item[1] for item in documents if item[0].name == name)

    headers: list[GroupItem] = []
    footers: list[GroupItem] = []
    for group in doc.groups:
        if not isinstance(group, GroupItem):
            continue
        if group.name == "page header":
            headers.append(group)
        elif group.name == "page footer":
            footers.append(group)

    assert len(headers) == 2, "Expected 2 different headers"
    assert len(footers) == 2, "Expected 2 different footers"

    assert len(headers[0].children) == 1, "First page header should have 1 paragraph"
    assert len(headers[1].children) == 2, "Second page header should have 2 paragraphs"

    assert len(footers[0].children) == 1, "First page footer should have 1 paragraph"
    assert len(footers[1].children) == 4, (
        "Second page footer should have 3 paragraphs and 1 picture"
    )


def test_handle_pictures(documents):
    """Test the function _handle_pictures."""

    name = "docx_grouped_images.docx"
    doc = next(item[1] for item in documents if item[0].name == name)

    assert len(doc.pictures) == 6
    assert isinstance(doc.pictures[0].parent.resolve(doc), GroupItem)
    assert doc.pictures[0].parent == doc.pictures[1].parent
    assert isinstance(doc.pictures[2].parent.resolve(doc), GroupItem)
    assert doc.pictures[2].parent == doc.pictures[3].parent
    assert isinstance(doc.pictures[4].parent.resolve(doc), SectionHeaderItem)
    assert doc.pictures[4].parent == doc.pictures[5].parent


def test_comments_extraction(documents):
    """Test the function _add_comments for extracting Word document comments."""

    name = "word_comments.docx"
    doc = next(item[1] for item in documents if item[0].name == name)

    # Find comment groups
    comment_groups: list[GroupItem] = []
    for group in doc.groups:
        if not isinstance(group, GroupItem):
            continue
        if group.name.startswith("comment-"):
            comment_groups.append(group)

    assert len(comment_groups) == 3, "Expected 3 comments in the document"

    # Collect all comment text content
    comment_texts = []
    for text_item in doc.texts:
        if hasattr(text_item, "content_layer") and text_item.content_layer == "notes":
            comment_texts.append(text_item.text)

    # Check that author info is included with new format
    assert any("author: John Reviewer (JR)" in text for text in comment_texts), (
        "Expected 'author: John Reviewer (JR)' in comments"
    )
    assert any("author: Jane Editor (JE)" in text for text in comment_texts), (
        "Expected 'author: Jane Editor (JE)' in comments"
    )

    # Check that comment text is included
    assert any("sample reviewer comment" in text for text in comment_texts), (
        "Expected comment text content"
    )
    assert any(
        "Another comment by a different reviewer" in text for text in comment_texts
    ), "Expected second comment text content"

    # Check content layer is NOTES
    for group in comment_groups:
        assert group.content_layer == "notes", (
            "Comments should be in NOTES content layer"
        )


@pytest.mark.parametrize(
    "style_label,expected_label,expected_level",
    [
        ("Heading 1", "Heading", 1),
        ("Heading 2", "Heading", 2),
        ("Heading 9", "Heading", 9),
        ("Heading 0", "Heading", 1),  # Custom style - level 0 should be clamped to 1
        ("1 Heading", "Heading", 1),  # Number before text
        ("0 Heading", "Heading", 1),  # Zero before text should be clamped to 1
    ],
)
def test_get_heading_and_level(docx_paths, style_label, expected_label, expected_level):
    """Test _get_heading_and_level handles edge cases like 'Heading 0' correctly."""
    # Create a backend instance using any existing docx file
    docx_path = docx_paths[0]
    in_doc = InputDocument(
        path_or_stream=docx_path,
        format=InputFormat.DOCX,
        backend=MsWordDocumentBackend,
    )
    backend = in_doc._backend

    label, level = backend._get_heading_and_level(style_label)
    assert label == expected_label, (
        f"Expected label '{expected_label}' for '{style_label}', got '{label}'"
    )
    assert level == expected_level, (
        f"Expected level {expected_level} for '{style_label}', got {level}"
    )


def test_get_outline_level_from_style():
    """Test that _get_outline_level_from_style correctly extracts outlineLvl.

    Uses word_sample.docx which has known heading paragraphs:
    - Paragraph 5: "Let's swim!" with Heading 1 style (outlineLvl=0 in XML)
    - Paragraph 15: "Let's eat" with Heading 2 style (outlineLvl=1 in XML)

    OOXML outlineLvl is 0-indexed, so our method should return outlineLvl + 1.
    """
    from docx import Document

    docx_path = Path("./tests/data/docx/word_sample.docx")
    in_doc = InputDocument(
        path_or_stream=docx_path,
        format=InputFormat.DOCX,
        backend=MsWordDocumentBackend,
    )
    backend = in_doc._backend
    doc = Document(docx_path)
    paragraphs = doc.paragraphs

    # Test Heading 1: outlineLvl=0 should return level 1
    heading1_para = paragraphs[5]
    assert heading1_para.text == "Let\u2019s swim!", "Test document structure changed"
    assert heading1_para.style.name == "Heading 1"
    assert backend._get_outline_level_from_style(heading1_para) == 1

    # Test Heading 2: outlineLvl=1 should return level 2
    heading2_para = paragraphs[15]
    assert heading2_para.text == "Let\u2019s eat", "Test document structure changed"
    assert heading2_para.style.name == "Heading 2"
    assert backend._get_outline_level_from_style(heading2_para) == 2

    # Test non-heading paragraph: should return None
    normal_para = paragraphs[0]  # First paragraph is not a heading
    assert "heading" not in normal_para.style.name.lower()
    assert backend._get_outline_level_from_style(normal_para) is None


@pytest.mark.parametrize(
    "style_label,expected_label,expected_level",
    [
        ("Normal", "Normal", None),  # Non-heading style
        ("Title", "Title", None),  # Non-heading style
        ("CustomStyle", "CustomStyle", None),  # Non-heading style
    ],
)
def test_get_heading_and_level_non_heading(
    docx_paths, style_label, expected_label, expected_level
):
    """Test _get_heading_and_level returns input unchanged for non-heading styles."""
    docx_path = docx_paths[0]
    in_doc = InputDocument(
        path_or_stream=docx_path,
        format=InputFormat.DOCX,
        backend=MsWordDocumentBackend,
    )
    backend = in_doc._backend

    label, level = backend._get_heading_and_level(style_label)
    assert label == expected_label
    assert level == expected_level


def test_external_image_references():
    """Test that .docx files with external image references convert without crashing.

    Docx files saved from web browsers often have images as external references
    (TargetMode="External") pointing to URLs or file:// paths rather than embedded
    in word/media/. Previously this caused a ValueError from python-docx:
    "target_part property on _Relationship is undefined when target mode is External"

    See: https://github.com/docling-project/docling/issues/3113
    """
    docx_path = Path("./tests/data/docx/docx_external_image.docx")
    assert docx_path.exists(), f"Test file not found: {docx_path}"

    converter = get_converter()

    with pytest.warns(UserWarning, match="Skipping external image reference"):
        conv_result = converter.convert(docx_path)

    doc = conv_result.document

    # Document should convert successfully (not crash)
    assert doc is not None

    # Text content should still be extracted even though the external image is skipped
    md = doc.export_to_markdown()
    assert "Test Document with External Image" in md
    assert "text before the image" in md
    assert "after the external image" in md


def test_inline_sdt_references(tmp_path):
    """Test that inline SDT citation blocks are preserved in DOCX paragraphs."""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def _append_citation(paragraph, text: str):
        sdt = OxmlElement("w:sdt")
        sdt_pr = OxmlElement("w:sdtPr")
        tag = OxmlElement("w:tag")
        tag.set(qn("w:val"), "MENDELEY_CITATION_v3_test")
        sdt_pr.append(tag)

        sdt_content = OxmlElement("w:sdtContent")
        run = OxmlElement("w:r")
        run_text = OxmlElement("w:t")
        run_text.text = text
        run.append(run_text)
        sdt_content.append(run)

        sdt.append(sdt_pr)
        sdt.append(sdt_content)
        paragraph._p.append(sdt)

    docx_path = tmp_path / "inline_sdt_reference.docx"
    doc = Document()

    first_paragraph = doc.add_paragraph()
    first_paragraph.add_run("Impact ")
    _append_citation(first_paragraph, "(Hagman G 1984)")
    first_paragraph.add_run(". After.")

    second_paragraph = doc.add_paragraph()
    _append_citation(second_paragraph, "(Standalone citation)")

    doc.save(docx_path)

    conv_result = get_converter().convert(docx_path)
    markdown = conv_result.document.export_to_markdown()

    assert "Impact (Hagman G 1984). After." in markdown
    assert "(Standalone citation)" in markdown


def test_list_counter_and_enum_marker(docx_paths):
    """Test list counter increment, sub-level reset, marker building, and sequence reset."""
    docx_path = docx_paths[0]
    in_doc = InputDocument(
        path_or_stream=docx_path,
        format=InputFormat.DOCX,
        backend=MsWordDocumentBackend,
    )
    backend = in_doc._backend

    # Basic increment
    assert backend._get_list_counter(1, 0) == 1
    assert backend._get_list_counter(1, 0) == 2
    assert backend._get_list_counter(1, 1) == 1
    assert backend._get_list_counter(1, 1) == 2
    assert backend._get_list_counter(1, 1) == 3

    # Advancing parent level resets sub-levels
    backend._get_list_counter(1, 2)  # (1,2) = 1
    backend._get_list_counter(1, 0)  # (1,0) = 3, resets lvl 1 and 2
    assert backend.list_counters[(1, 1)] == 0
    assert backend.list_counters[(1, 2)] == 0
    assert backend._get_list_counter(1, 1) == 1  # restarts from 1

    # Hierarchical enum markers
    backend.list_counters[(1, 0)] = 2
    backend.list_counters[(1, 1)] = 3
    backend.list_counters[(1, 2)] = 1
    assert backend._build_enum_marker(1, 0) == "2."
    assert backend._build_enum_marker(1, 1) == "2.3."
    assert backend._build_enum_marker(1, 2) == "2.3.1."
    assert backend._build_enum_marker(99, 0) == "1."  # missing counter defaults to 1

    # Reset sequence for a specific numid
    backend._get_list_counter(2, 0)  # (2,0) = 1
    backend._reset_list_counters_for_new_sequence(1)
    assert backend.list_counters[(1, 0)] == 0
    assert backend.list_counters[(1, 1)] == 0
    assert backend.list_counters[(2, 0)] == 1  # unaffected


def test_handle_equations_in_text_returns_original_text_on_mismatch(
    backend, monkeypatch
):
    element = etree.Element("p")
    run = etree.SubElement(element, "r")
    text_elem = etree.SubElement(run, "t")
    text_elem.text = "alpha"
    etree.SubElement(element, "oMath")

    monkeypatch.setattr(msword_backend_module, "oMath2Latex", lambda _: "x")

    text, equations = backend._handle_equations_in_text(element=element, text="beta")

    assert text == "beta"
    assert equations == []


def test_handle_equations_in_text_skips_empty_substrings(backend, monkeypatch):
    equation = backend.equation_bookends.format(EQ="x")

    element = etree.Element("p")
    empty_run = etree.SubElement(element, "r")
    empty_text = etree.SubElement(empty_run, "t")
    empty_text.text = ""
    etree.SubElement(element, "oMath")
    tail_run = etree.SubElement(element, "r")
    tail_text = etree.SubElement(tail_run, "t")
    tail_text.text = "tail"

    monkeypatch.setattr(msword_backend_module, "oMath2Latex", lambda _: "x")

    text, equations = backend._handle_equations_in_text(element=element, text="tail")

    assert equations == [equation]
    assert text == f"{equation}tail"


def test_handle_text_elements_returns_empty_refs_when_text_is_none(
    backend, monkeypatch
):
    element = backend.docx_obj.paragraphs[0]._element

    monkeypatch.setattr(
        backend, "_handle_equations_in_text", lambda element, text: (None, [])
    )

    refs = backend._handle_text_elements(element, DoclingDocument(name="test"))

    assert refs == []


def test_handle_text_elements_heading_defaults_to_non_numbered_when_style_missing(
    backend, monkeypatch
):
    captured: dict[str, tuple[int, str, bool]] = {}

    class FakeParagraph:
        def __init__(self, element, docx_obj):
            self.text = "Heading text"
            self.style = SimpleNamespace()

    monkeypatch.setattr(msword_backend_module, "Paragraph", FakeParagraph)
    monkeypatch.setattr(backend, "_get_paragraph_elements", lambda paragraph: [])
    monkeypatch.setattr(
        backend, "_handle_equations_in_text", lambda element, text: (text, [])
    )
    monkeypatch.setattr(backend, "_get_comment_ids_for_element", lambda element: [])
    monkeypatch.setattr(
        backend, "_get_label_and_level", lambda paragraph: ("Heading", 1)
    )
    monkeypatch.setattr(backend, "_get_numId_and_ilvl", lambda paragraph: (None, None))

    def fake_add_heading(doc, level, text, is_numbered_style):
        captured["heading"] = (level, text, is_numbered_style)
        return []

    monkeypatch.setattr(backend, "_add_heading", fake_add_heading)

    refs = backend._handle_text_elements(object(), DoclingDocument(name="test"))

    assert refs == []
    assert captured["heading"] == (1, "Heading text", False)


def test_handle_text_elements_inline_equations_stop_when_text_is_consumed(
    backend, monkeypatch
):
    equation_one = backend.equation_bookends.format(EQ="a")
    equation_two = backend.equation_bookends.format(EQ="b")

    class FakeParagraph:
        def __init__(self, element, docx_obj):
            self.text = "inline eq"
            self.style = SimpleNamespace()

    monkeypatch.setattr(msword_backend_module, "Paragraph", FakeParagraph)
    monkeypatch.setattr(backend, "_get_paragraph_elements", lambda paragraph: [])
    monkeypatch.setattr(
        backend,
        "_handle_equations_in_text",
        lambda element, text: (equation_one, [equation_one, equation_two]),
    )
    monkeypatch.setattr(backend, "_get_comment_ids_for_element", lambda element: [])
    monkeypatch.setattr(
        backend, "_get_label_and_level", lambda paragraph: ("Normal", None)
    )
    monkeypatch.setattr(backend, "_get_numId_and_ilvl", lambda paragraph: (None, None))
    monkeypatch.setattr(backend, "_prev_numid", lambda: None)
    monkeypatch.setattr(backend, "_get_level", lambda: 1)
    backend.parents[0] = None

    refs = backend._handle_text_elements(object(), DoclingDocument(name="test"))

    assert len(refs) == 2


def test_checkbox_detection_and_parsing(documents):
    """Test that checkboxes in DOCX files are correctly detected and parsed."""
    name = "docx_checkboxes.docx"
    doc = next((item[1] for item in documents if item[0].name == name), None)

    if doc is None:
        pytest.skip(f"Test file not found: {name}")

    checkbox_items = [
        item
        for item in doc.texts
        if item.label
        in (DocItemLabel.CHECKBOX_SELECTED, DocItemLabel.CHECKBOX_UNSELECTED)
    ]

    assert len(checkbox_items) > 0, "No checkboxes found in the document"

    # Verify we have both selected and unselected checkboxes
    selected = [
        item for item in checkbox_items if item.label == DocItemLabel.CHECKBOX_SELECTED
    ]
    unselected = [
        item
        for item in checkbox_items
        if item.label == DocItemLabel.CHECKBOX_UNSELECTED
    ]

    assert len(selected) > 0, "No selected checkboxes found"
    assert len(unselected) > 0, "No unselected checkboxes found"

    checkbox_texts = [item.text for item in checkbox_items]
    assert any("Design" in text for text in checkbox_texts), (
        "Expected checkbox text not found"
    )
    assert any("Implementation" in text for text in checkbox_texts), (
        "Expected checkbox text not found"
    )
    assert any("Documentation" in text for text in checkbox_texts), (
        "Expected checkbox text not found"
    )


def test_checkbox_labels_in_tables(documents):
    """Test that checkboxes in table cells are correctly parsed."""
    name = "docx_checkboxes.docx"
    doc = next((item[1] for item in documents if item[0].name == name), None)

    if doc is None:
        pytest.skip(f"Test file not found: {name}")

    checkbox_items = [
        item
        for item in doc.texts
        if item.label
        in (DocItemLabel.CHECKBOX_SELECTED, DocItemLabel.CHECKBOX_UNSELECTED)
    ]

    food_items = [
        "Orange juice",
        "Tea",
        "Coffee",
        "Milk",
        "Water",
        "Scramble eggs",
        "Porridge",
        "Bread",
        "Croissant",
    ]

    found_food_checkboxes = [
        item for item in checkbox_items if any(food in item.text for food in food_items)
    ]

    assert len(found_food_checkboxes) > 0, "No checkboxes found in table cells"
