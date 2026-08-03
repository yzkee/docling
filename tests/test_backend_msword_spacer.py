"""Tests for Word list handling around blank spacer paragraphs.

Kept separate from ``test_backend_msword.py`` so that file stays under the
repository's per-file line limit.
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docling.datamodel.base_models import InputFormat
from docling.document_converter import DocumentConverter


def _make_multilevel_numbering(doc, abstract_id: str, num_id: str, levels: int = 3):
    """Register a multi-level decimal numbering definition in a docx."""
    numbering = doc.part.numbering_part.element

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), abstract_id)
    for ilvl in range(levels):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        numfmt = OxmlElement("w:numFmt")
        numfmt.set(qn("w:val"), "decimal")
        lvl.append(numfmt)
        lvltext = OxmlElement("w:lvlText")
        lvltext.set(qn("w:val"), ".".join(f"%{i + 1}" for i in range(ilvl + 1)))
        lvl.append(lvltext)
        abstract_num.append(lvl)
    numbering.append(abstract_num)

    num_elem = OxmlElement("w:num")
    num_elem.set(qn("w:numId"), num_id)
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_id)
    num_elem.append(abstract_ref)
    numbering.append(num_elem)


def _add_numbered_paragraph(doc, text: str, num_id: str, ilvl: int):
    paragraph = doc.add_paragraph(text)
    num_pr = OxmlElement("w:numPr")
    ilvl_elem = OxmlElement("w:ilvl")
    ilvl_elem.set(qn("w:val"), str(ilvl))
    num_pr.append(ilvl_elem)
    num_id_elem = OxmlElement("w:numId")
    num_id_elem.set(qn("w:val"), num_id)
    num_pr.append(num_id_elem)
    paragraph._p.get_or_add_pPr().append(num_pr)
    return paragraph


def test_empty_paragraph_between_list_items_keeps_body_text_in_place(tmp_path):
    """A blank spacer paragraph must not strand the body text after it.

    Authors commonly press Enter between list items for vertical spacing. Such
    an empty paragraph closes the list without clearing the cached list group,
    so the next item used to re-open that group -- which sits before the body
    text in between -- and the text ended up after the whole list instead of
    where the author put it.
    """

    converter = DocumentConverter(allowed_formats=[InputFormat.DOCX])

    def build(with_spacer: bool) -> list[str]:
        doc = Document()
        _make_multilevel_numbering(doc, abstract_id="700", num_id="701")

        _add_numbered_paragraph(doc, "First section", "701", 0)
        _add_numbered_paragraph(doc, "Sub one", "701", 1)
        if with_spacer:
            doc.add_paragraph("")
        doc.add_paragraph("Prose that belongs under Sub one.")
        _add_numbered_paragraph(doc, "Sub two", "701", 1)
        _add_numbered_paragraph(doc, "Second section", "701", 0)

        name = "with_spacer" if with_spacer else "without_spacer"
        docx_path = tmp_path / f"{name}.docx"
        doc.save(str(docx_path))
        markdown = converter.convert(docx_path).document.export_to_markdown()
        return [line for line in markdown.splitlines() if line.strip()]

    lines = build(with_spacer=True)
    # The spacer changes nothing about where the content ends up.
    assert lines == build(with_spacer=False)

    prose = lines.index("Prose that belongs under Sub one.")
    sub_one = next(i for i, line in enumerate(lines) if line.endswith("Sub one"))
    sub_two = next(i for i, line in enumerate(lines) if line.endswith("Sub two"))
    assert sub_one < prose < sub_two
