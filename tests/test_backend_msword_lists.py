# SPDX-FileCopyrightText: The Docling Contributors
# SPDX-License-Identifier: MIT

"""Tests for Word list numbering behaviour.

Kept separate from ``test_backend_msword.py`` so that file stays under the
repository's per-file line limit.
"""

from io import BytesIO

from docling_core.types.doc import DoclingDocument, DocumentOrigin, ListGroup, ListItem
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docling.backend.msword_backend import MsWordDocumentBackend
from docling.datamodel.base_models import InputFormat
from docling.datamodel.document import InputDocument


def test_ordered_list_resumes_numbering_after_intervening_list(tmp_path):
    """An ordered list interrupted by a bullet list must keep counting.

    Word numbers continuously per ``w:numId``, so items sharing one numId are a
    single list even when other content (including a list with a different
    numId) sits between them. Docling used to reset the counter whenever the
    numId changed, so the resumed items restarted at 1.
    """

    doc = Document()
    numbering = doc.part.numbering_part.element

    def add_numbering(abstract_id: str, num_id: str, num_fmt: str, lvl_text: str):
        abstract_num = OxmlElement("w:abstractNum")
        abstract_num.set(qn("w:abstractNumId"), abstract_id)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_fmt)
        lvl.append(fmt)
        text = OxmlElement("w:lvlText")
        text.set(qn("w:val"), lvl_text)
        lvl.append(text)
        abstract_num.append(lvl)
        numbering.append(abstract_num)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), num_id)
        ref = OxmlElement("w:abstractNumId")
        ref.set(qn("w:val"), abstract_id)
        num.append(ref)
        numbering.append(num)

    add_numbering("300", "301", "decimal", "%1.")
    add_numbering("400", "401", "bullet", "•")

    def add_item(text: str, num_id: str):
        paragraph = doc.add_paragraph(text, style="List Paragraph")
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_pr.append(ilvl)
        num_id_elem = OxmlElement("w:numId")
        num_id_elem.set(qn("w:val"), num_id)
        num_pr.append(num_id_elem)
        paragraph._element.get_or_add_pPr().append(num_pr)

    add_item("First ordered item", "301")
    add_item("Second ordered item", "301")
    add_item("bullet one", "401")
    add_item("bullet two", "401")
    add_item("Third ordered item", "301")

    docx_path = tmp_path / "resumed_ordered_list.docx"
    doc.save(str(docx_path))

    in_doc = InputDocument(
        path_or_stream=docx_path,
        format=InputFormat.DOCX,
        backend=MsWordDocumentBackend,
        filename=docx_path.name,
    )
    converted = MsWordDocumentBackend(in_doc=in_doc, path_or_stream=docx_path).convert()

    markers = [
        (item.text, item.marker)
        for item, _ in converted.iterate_items()
        if isinstance(item, ListItem)
    ]

    assert markers == [
        ("First ordered item", "1."),
        ("Second ordered item", "2."),
        ("bullet one", ""),
        ("bullet two", ""),
        ("Third ordered item", "3."),
    ]


def test_manage_list_structure_no_keyerror_when_use_level_exceeds_parents(tmp_path):
    """_manage_list_structure must not raise KeyError when use_level exceeds parents.

    The pathological state is:
      - parents has keys 0..11, with key 0 set to None (cleared) and others
        holding NodeItems from headings / earlier lists.
      - level_at_new_list is 11 (set by the previous list item of the same numId).
      - A second item for the same numId arrives with ilevel=2, which triggers
        the "New list sequence" branch and computes use_level = 11 + 2 = 13.
      - parents.get(12) is not a key at all → previously raised KeyError: 12.
    """

    docx_io = _make_empty_docx()
    docx_path = tmp_path / "empty.docx"
    docx_path.write_bytes(docx_io.getvalue())

    in_doc = InputDocument(
        path_or_stream=docx_path,
        format=InputFormat.DOCX,
        backend=MsWordDocumentBackend,
        filename=docx_path.name,
    )
    backend = MsWordDocumentBackend(in_doc=in_doc, path_or_stream=docx_path)

    out_doc = DoclingDocument(
        name="test",
        origin=DocumentOrigin(
            filename="test.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            binary_hash="0",
        ),
    )

    body_node = out_doc.body
    first_list_gr = out_doc.add_list_group(name="list", parent=body_node)
    backend.parents = dict.fromkeys(range(12), body_node)
    backend.parents[0] = None  # gap that makes _get_level() return 0
    backend.parents[11] = first_list_gr

    # Simulate history: the immediately previous item was numid=27, ilevel=2.
    backend.history = {
        "names": [None, "list"],
        "levels": [None, 11],
        "numids": [None, 27],
        "indents": [None, 2],
    }
    backend.level_at_new_list = 11
    backend.last_numid = 27

    # This must not raise KeyError.
    elem_ref, use_level = backend._manage_list_structure(
        doc=out_doc, numid=27, ilevel=2
    )

    assert isinstance(backend.parents.get(use_level), ListGroup)


def test_manage_list_structure_no_keyerror_open_indented_list_exceeds_parents(tmp_path):
    """_manage_list_structure must not raise KeyError in the "Open indented list" branch.

    The pathological state is:
      - parents has keys 0..11, with keys 0..10 holding body nodes and key 11
        holding a ListGroup.
      - level_at_new_list is 11, prev_indent is 2, ilevel is 4.
      - The "Open indented list" loop runs for i in range(14, 16), accessing
        parents[i - 1] where i - 1 = 13 is not a key in parents.
    """
    docx_path = tmp_path / "empty.docx"
    docx_path.write_bytes(_make_empty_docx().getvalue())

    in_doc = InputDocument(
        path_or_stream=docx_path,
        format=InputFormat.DOCX,
        backend=MsWordDocumentBackend,
        filename=docx_path.name,
    )
    backend = MsWordDocumentBackend(in_doc=in_doc, path_or_stream=docx_path)

    out_doc = DoclingDocument(
        name="test",
        origin=DocumentOrigin(
            filename="test.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            binary_hash="0",
        ),
    )

    body_node = out_doc.body
    first_list_gr = out_doc.add_list_group(name="list", parent=body_node)
    backend.parents = dict.fromkeys(range(12), body_node)
    backend.parents[11] = first_list_gr

    # Simulate history: same numId, previous item was at ilevel=2.
    backend.history = {
        "names": [None, "list"],
        "levels": [None, 11],
        "numids": [None, 27],
        "indents": [None, 2],
    }
    backend.level_at_new_list = 11
    backend.last_numid = 27

    elem_ref, use_level = backend._manage_list_structure(
        doc=out_doc, numid=27, ilevel=4
    )

    assert isinstance(backend.parents.get(use_level), ListGroup)


def _make_empty_docx():
    """Return an in-memory .docx with no content."""

    buf = BytesIO()
    Document().save(buf)
    buf.seek(0)
    return buf
